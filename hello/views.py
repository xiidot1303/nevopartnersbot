from django.shortcuts import render, redirect
from telegram import ext
from .models import *
from .forms import *
from django.views.generic.edit import UpdateView, CreateView, DeleteView
from django.views.generic.detail import DetailView
from django.urls import reverse
from django.http import FileResponse
from django.http import HttpResponse, JsonResponse, FileResponse
from django.views.decorators.http import require_POST
from django.template.loader import render_to_string
from django import forms
from django.conf import settings
from django.core.files.storage import FileSystemStorage
from telegram.ext import Updater, CommandHandler, MessageHandler, Filters, CallbackQueryHandler
from telegram import InlineKeyboardButton, InlineKeyboardMarkup
import requests
import json
#from .bot import updater
from django.views import View
from django.contrib.auth.mixins import LoginRequiredMixin
from django.http import HttpResponse
import telegram
from django.contrib.auth.models import User, Permission
from os import getenv
from datetime import datetime
from django.contrib.auth.mixins import LoginRequiredMixin, PermissionRequiredMixin
from django.contrib.auth.decorators import login_required, permission_required
from django.views.decorators.http import require_GET
from django.views.decorators.csrf import csrf_exempt
from django.db.models import Q
from telegram import Update
from .dispatcher import dp

from docx import Document
from docx.shared import Inches, Pt

from django.forms import modelformset_factory, formset_factory
from dotenv import load_dotenv
import os

basedir = os.path.abspath(os.path.dirname(''))
load_dotenv(os.path.join(basedir, '.env'))
TOKEN = os.environ.get('TOKEN')


BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))        

@require_GET
def robots_txt(request):
    lines = [
        "User-Agent: *",
        "Disallow: /private/",
        "Disallow: /junk/",
    ]
    return HttpResponse("\n".join(lines), content_type="text/plain")



@csrf_exempt
def bot_webhook(request):
    #try:
    #   
    update = Update.de_json(json.loads(request.body.decode('utf-8')), dp.bot)
    dp.process_update(update)
    return HttpResponse('Bot started!')
    #else:

    #    return render(request, 'bot/startbot.html')
    #except Exception as e:
#
    #    return HttpResponse(e)



class StartBotView(View):
    def get(self, request, *args, **kwargs):
        
        updater.start_webhook(listen='127.0.0.1', port=12338, url_path='bot-secret')
        updater.bot.set_webhook(url='https://{host}/{path}'.format(host='nevopartners.elite-house.uz', path='bot-secret'))
        return HttpResponse('Bot started!')




@login_required
def Sendmessage(request, ps, issent):
    if request.method == 'POST':
        bbf = SendmessageForm(request.POST)
        if bbf.is_valid():
            sendmessage.objects.create(admin = request.user, select = bbf.cleaned_data['select'], message=bbf.cleaned_data['message'], user=bbf.cleaned_data['select'])
            if bbf.cleaned_data['select'] == 'Все':
                msg = bbf.cleaned_data['message']
                users = subscribersbot.objects.all()
                for u in users:
                    my_token = TOKEN
                    bot = telegram.Bot(token=my_token)
                    try:
                        bot.sendMessage(chat_id=u.user_id, text='Сообщения от админстратора:\n'+msg)
                    except:
                        fewfwe = 0

            else:
                pseu = bbf.cleaned_data['select']
                msg = bbf.cleaned_data['message']
                login = Profile.objects.get(pseudonym=pseu).login
                user_id = subscribersbot.objects.get(login=login).user_id
                my_token = TOKEN


                bot = telegram.Bot(token=my_token)
                
                bot.sendMessage(chat_id=user_id, text='Сообщения от админстратора:\n'+msg)

                return redirect('sendmessage', permanent=True, ps=ps, issent='yes')
            
            
            return redirect('profiles', permanent=False)
            
        else:
            profiles = Profile.objects.all()
            try:
                current_profil = Profile.objects.get(pseudonym=ps).pk
            except:
                current_profil = None
            if ps == 'all':
                story = sendmessage.objects.all()
            else:    
                story = sendmessage.objects.filter(user=ps).order_by('published')
            context = {'form': bbf, 'ps': ps, 'profiles': profiles, 'issent': issent, 'current': current_profil, 'stories': story}
            return render(request, 'bot/sendmessage.html', context)

    else:
        bbf = SendmessageForm()
        profiles = Profile.objects.all()
        try:
            current_profil = Profile.objects.get(pseudonym=ps).pk
        except:
            current_profil = None
        if ps == 'all':
            story = sendmessage.objects.all().order_by('published')
        else:    
            story = sendmessage.objects.filter(user=ps).order_by('published')
        context = {'form': bbf, 'ps': ps, 'profiles': profiles, 'issent': issent, 'current': current_profil, 'stories': story}
        return render(request, 'bot/sendmessage.html', context)    



@login_required
def folder(request):
    profiles = Profile.objects.all()
    prefix = []
    pks = []
    for p in profiles:
        if Account.objects.filter(pseudonym=p.pseudonym):
            prefix.append(p.prefix)
            pks.append(p.pk)
    card_number = []
    for c in profiles:
        if Account.objects.filter(pseudonym=c.pseudonym):
            card_number.append(c.card_number)
    table = []
    minx = int(Account.objects.values_list('year').order_by('-year').first()[0])
    maxx = int(Account.objects.values_list('year').order_by('-year').last()[0])
    for i in profiles:
        a = Account.objects.filter(pseudonym=i.pseudonym)
        if a:
            table.append(a)
    # months = ['Январь', 'Февраль', 'Март', 'Апрель', 'Май', 'Июнь', 'Июль', 'Август', 'Сентябрь', 'Октябрь', 'Ноябрь', 'Декабрь']
    months = ['Декабрь', 'Ноябрь', 'Октябрь', 'Сентябрь', 'Август', 'Июль', 'Июнь', 'Май', 'Апрель', 'Март', 'Февраль', 'Январь']
    min_month = Account.objects.filter(year=minx).values_list('month').order_by('-month').first()[0]

    first_year = minx
    first_months = months[-(min_month):]
    first_months_ids = list(range(1, min_month+1))
    first_months_ids.reverse()
    print(first_months_ids)

    ids = [1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12]
    ids = [12, 11, 10, 9, 8, 7, 6, 5, 4, 3, 2, 1]

    years = list(range(maxx, minx))
    years.reverse()
    context = {'prefix': prefix, 'card_number': card_number, 'profiles': profiles, 'ps': 'Все', 'allprofiles': profiles, 'path': 'Все', 'table': table, 
    'ids': ids, 'months': months, 'years': years, 'first_year': first_year, 'first_months': first_months, 'first_months_ids': first_months_ids, 'pks': pks}
    return render(request, 'bot/folder.html', context)
@login_required
def sortfolder(request, ps):
    profiles = Profile.objects.all()
    if ps != 'Все':
        profile = Profile.objects.get(pseudonym=ps)
        profiles = [profile]
    allprofiles = Profile.objects.all()
    # -----
    prefix = []
    for p in profiles:
        if Account.objects.filter(pseudonym=p.pseudonym):
            prefix.append(p.prefix)
            pks.append(p.pk)
    card_number = []
    for c in profiles:
        if Account.objects.filter(pseudonym=c.pseudonym):
            card_number.append(c.card_number)
    table = []
    for i in profiles:
        a = Account.objects.filter(pseudonym=i.pseudonym)
        if a:
            table.append(a)
    months = ['Январь', 'Февраль', 'Март', 'Апрель', 'Май', 'Июнь', 'Июль', 'Август', 'Сентябрь', 'Октябрь', 'Ноябрь', 'Декабрь']
    
    ids = [1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12]
    #========
    if ps == 'Все':
        minx = int(Account.objects.values_list('year').order_by('year').first()[0])
        maxx = int(Account.objects.values_list('year').order_by('year').last()[0])
        min_month = Account.objects.filter(year=minx).values_list('month').order_by('month').first()[0]
    else:
        minx = int(Account.objects.filter(pseudonym=ps).values_list('year').order_by('year').first()[0])
        maxx = int(Account.objects.filter(pseudonym=ps).values_list('year').order_by('year').last()[0])
        min_month = Account.objects.filter(year=minx, pseudonym=ps).values_list('month').order_by('month').first()[0]
    first_year = minx
    first_months = months[(min_month-1):]
    first_months_ids = range(min_month, 13)
    context = {'prefix': prefix, 'card_number': card_number, 'profiles': profiles, 'ps': ps, 'allprofiles': allprofiles, 'path': ps, 'table': table, 'ids': ids, 'months': months, 'years': range(minx+1, maxx+1), 
    'first_year': first_year, 'first_months': first_months, 'first_months_ids': first_months_ids, 'pks': pks}
    
    
    return render(request, 'bot/folder.html', context)



@login_required
def accounts(request, pse):
    ps = Profile.objects.get(prefix=pse).pseudonym
    accs = Account.objects.filter(pseudonym=ps).order_by('year', 'month_id')
    context = {'accounts': accs}
    return render(request, 'bot/account.html', context)
@login_required
def allaccounts(request):
    obj = Account.objects.all().order_by('-published', 'pseudonym')
    context = {'accounts': obj, 'ps': 'Все', 'status':'Все', 'year':'Все', 'month':'1'}
    context['years'] = ['Все', '2016', '2017', '2018', '2019', '2020', '2021', '2022', '2023', '2024', '2025', '2026', '2027']
    context['months'] = ['Все', 'Январь', 'Февраль', 'Март', 'Апрель', 'Май', 'Июнь', 'Июль', 'Август', 'Сентябрь', 'Октябрь', 'Ноябрь', 'Декабрь']
    context['statuses'] = ['Все', 'В работе', 'В ожидании', 'Оплачен', 'Отклонена', 'Перенесена', 'Нет']
    context['profiles'] = Profile.objects.all().order_by('pseudonym')
    context['path'] = 'Все'
    return render(request, 'bot/allaccount.html', context)
@login_required
def sortallaccounts(request, ps, status, year, month):
    obj = Account.objects.all().order_by('-published', 'pseudonym')
    if ps != 'Все':
        obj = obj.filter(pseudonym=ps)
    if status != 'Все':
        obj = obj.filter(status__s=status)
    try:
        obj = obj.filter(year=int(year))
    except:
        fwrferwfrwe=0
    
    if month != '1':
        obj = obj.filter(month=int(month)-1)
    

    context = {'accounts': obj, 'ps': ps, 'status': status, 'year': year, 'month': month}
    context['years'] = ['Все', '2016', '2017', '2018', '2019', '2020', '2021', '2022', '2023', '2024', '2025', '2026', '2027']
    context['months'] = ['Все', 'Январь', 'Февраль', 'Март', 'Апрель', 'Май', 'Июнь', 'Июль', 'Август', 'Сентябрь', 'Октябрь', 'Ноябрь', 'Декабрь']
    context['statuses'] = ['Все', 'В работе', 'В ожидании', 'Оплачен', 'Отклонена', 'Перенесена', 'Нет']
    context['profiles'] = Profile.objects.all().order_by('pseudonym')
    context['path'] = ps
    return render(request, 'bot/allaccount.html', context)

class AccCreateView(LoginRequiredMixin, CreateView, PermissionRequiredMixin):
    template_name = 'bot/createacc.html'
    form_class = AccForm
    success_url = '/detail/{id}'
    permanent = True
    permission_required = 'hello.add_account'
    def get_context_data(self, *args, **kwargs):
        context=super().get_context_data(*args, **kwargs)
        

        ps = []
        for i in Profile.objects.all().order_by('pseudonym'):
            ps.append(i.pseudonym)
        context['profiles'] = ps
        
        context['cards'] = card.objects.all()
        context['allprofiles'] = Profile.objects.all().order_by('pseudonym')
        return context

class AccDeleteView(LoginRequiredMixin, DeleteView, PermissionRequiredMixin):
    model = Account
    permission_required = 'hello.delete_account'
    def get_success_url(self, *args, **kwargs):
        con = super().get_context_data(*args, **kwargs)
        obj = con['object']
        url = '/account/{}/'.format(obj.pseudonym)
        return url


class ProfDetailView(LoginRequiredMixin, DetailView):
    model = Profile
    
    def get_context_data(self, *args, **kwargs):
        context=super().get_context_data(*args, **kwargs)
        obj = context['object']
        if card.objects.filter(pseudonym=obj.pseudonym):   #checking that is creating or editing

            context['h1'] = 'Данные были успешно изменены'
            try:
                card.objects.get_or_create(card_number=obj.card_number, pseudonym=obj.pseudonym, type_payment=obj.type_payment, valute_card=obj.valute_card, owner_card=obj.owner_card)  
            except:
                sasa = 0
            if not subscribersbot.objects.filter(login=obj.login):
                for i in subscribersbot.objects.all():
                    if not Profile.objects.filter(login=i.login):
                        if i.parol:
                            token = TOKEN
                            user_id = i.user_id
                            bot = telegram.Bot(token=token)

                            message = 'Ваши данные для входа сменены. Для получение новых данных обратитесь Вашему менеджеру по работе с партнерами или по почты support@nevo.uz\n\nнажмите /start, чтобы перезапустить бота'
                            bot.sendMessage(chat_id=user_id, text=message)
                            changing.objects.get(user_id=i.user_id).delete()
                            typing.objects.get(user_id=i.user_id).delete()

                            try:
                                storage.objects.get(user_id=i.user_id).delete()
                            except:
                                qpwomsqd = 0
                            i.delete()
                        else:
                            t=typing.objects.get(user_id=i.user_id)
                            t.login=True
                            t.parol=False
                            t.save()
                            i.delete()
            else:
                if obj.parol != subscribersbot.objects.get(login=obj.login).parol and subscribersbot.objects.get(login=obj.login).parol:
                    token = TOKEN
                    i = subscribersbot.objects.get(login=obj.login)
                    user_id = i.user_id
                    bot = telegram.Bot(token=token)
                    
                    message = 'Ваши данные для входа сменены. Для получение новых данных обратитесь Вашему менеджеру по работе с партнерами или по почты support@nevo.uz\n\nнажмите /start, чтобы перезапустить бота'
                    try:
                        bot.sendMessage(chat_id=user_id, text=message)
                    except:
                        asasa = 0 # do nothing
                    changing.objects.get(user_id=i.user_id).delete()
                    typing.objects.get(user_id=i.user_id).delete()
                    
                    try:
                        storage.objects.get(user_id=i.user_id).delete()
                    except:
                        qpwomsqd = 0
                    i.delete()

        else:
            context['h1'] = 'Новый пользователь успешно добавлен'
            # prof pre values also will create at line 388
            card.objects.create(card_number=obj.card_number, pseudonym=obj.pseudonym, type_payment=obj.type_payment, valute_card=obj.valute_card, owner_card=obj.owner_card)
        
        now = Profile.objects.filter(pk=obj.pk).values()[0]
        c = 0
        try:
            pre = prof_pre_value.objects.filter(pk=obj.pk).values()[0]
            text = ''
            for i in pre:
                if i == 'id':
                    asddw = 0
                elif i == 'comment':
                    dcwcwe = 0
                elif pre[i] != now[i]:
                    text += i + ':\n' + str(pre[i]) + ' >> ' + str(now[i]) + ';    '
            if text:
                stories.objects.create(obj_id=obj.pk, admin=self.request.user, obj='Профил', text=text)
                
            
        except:
            stories.objects.create(obj_id=obj.pk, admin=self.request.user, obj='Профил', text="Создан")

        try:
            pre = prof_pre_value.objects.get(pk=obj.pk)
            # check that Accounts, Audios, Videos, cards didn't lost
            if obj.pseudonym != pre.pseudonym:
                for i in Account.objects.filter(pseudonym=pre.pseudonym):
                    i.pseudonym = obj.pseudonym
                    i.save()
                for i in Audio.objects.filter(pseudonym=pre.pseudonym):
                    i.pseudonym = obj.pseudonym
                    i.save()
                for i in Video.objects.filter(pseudonym=pre.pseudonym):
                    i.pseudonym = obj.pseudonym
                    i.save()
                for i in card.objects.filter(pseudonym=pre.pseudonym):
                    i.pseudonym = obj.pseudonym
                    i.save()
                os.system('mv {} {}'.format(os.path.join(BASE_DIR, 'files/app/{}'.format(pre.pseudonym)), os.path.join(BASE_DIR, 'files/app/{}'.format(obj.pseudonym))))

            pre.login = obj.login
            pre.parol = obj.parol
            pre.pseudonym = obj.pseudonym
            pre.number = obj.number
            pre.code = obj.code
            pre.name = obj.name
            pre.prefix = obj.prefix
            pre.type_payment = obj.type_payment
            pre.valute_card = obj.valute_card
            pre.card_number = obj.card_number
            pre.date_card = obj.date_card
            pre.owner_card = obj.owner_card
            pre.published = obj.published
            pre.date_birthday = obj.date_birthday
            pre.number_passport = obj.number_passport
            pre.date_passport = obj.date_passport
            pre.who_gave = obj.who_gave
            pre.place_registration = obj.place_registration
            pre.email = obj.email
            pre.bank = obj.bank
            pre.reward = obj.reward
            pre.passport_main_st = obj.passport_main_st
            pre.passport_st_registration = obj.passport_st_registration
            pre.signed_contract = obj.signed_contract
            pre.inn = obj.inn
            pre.type_document = obj.type_document
            pre.photo = obj.photo

            pre.save()
        except:
            # create , when new profile is added. 
            prof_pre_value.objects.create(pk=obj.pk, inn=obj.inn, published = obj.published, login = obj.login, parol = obj.parol, pseudonym = obj.pseudonym, number = obj.number, code = obj.code, name = obj.name, prefix = obj.prefix, type_payment = obj.type_payment, valute_card = obj.valute_card, card_number = obj.card_number, date_card = obj.date_card, owner_card = obj.owner_card, date_birthday = obj.date_birthday, number_passport = obj.number_passport, date_passport = obj.date_passport, who_gave = obj.who_gave, place_registration = obj.place_registration, email = obj.email, bank = obj.bank, reward = obj.reward, passport_main_st = obj.passport_main_st, passport_st_registration = obj.passport_st_registration, signed_contract = obj.signed_contract)

        return context

class ProfEditView(LoginRequiredMixin, UpdateView, PermissionRequiredMixin):
    model = Profile
    form_class = ProfForm
    success_url = '/new/{id}'
    permission_required = 'hello.change_profile'
    def get_context_data(self, *args, **kwargs):
        con = super().get_context_data(*args, **kwargs)
        obj = con['object']
        con['stories'] = stories.objects.filter(obj_id=obj.pk, obj='Профил').order_by('-published')
        
        return con
class AccDetailView(LoginRequiredMixin, DetailView):
    model = Account
    
    def get_context_data(self, *args, **kwargs):
        context=super().get_context_data(*args, **kwargs)
        obj = context['object']
        year = obj.year
        summa = obj.summa
        year = obj.year
        mo = obj.month.id
        months = ['Январь', 'Февраль', 'Март', 'Апрель', 'Май', 'Июнь', 'Июль', 'Август', 'Сентябрь', 'Октябрь', 'Ноябрь', 'Декабрь']
        month = months[mo-1]
        #__bot
        if obj.notification == 'Вкл':
            try:
                prof = Profile.objects.get(pseudonym=obj.pseudonym)
                fio = prof.name
                try:
                    c = card.objects.get(card_number=str(obj.card_number))
                    type_payment = c.type_payment
                    valute_card = c.valute_card
                    card_number = str(c.card_number)
                    owner_card = c.owner_card

                except:
                    type_payment = ''
                    valute_card = ''
                    card_number = ''
                    owner_card = ''


                l = prof.login
                user_id = subscribersbot.objects.get(login = l).user_id
                doc_send = True
                message = '{}, здравствуйте!\nМы подготовили выплату Вашего дохода на YouTube и просим проверить реквизиты:\n\nСпособ оплаты: {}\nВалюта карты: {}\nНомер карты: {}\nВладелец карты: {}\n\nПри отсутствии претензий или запроса по смены реквизитов в течение 3-х рабочих дней, отчет считается подтвержденным и выплата производится до конца месяца.'.format(fio, type_payment, valute_card, card_number, owner_card)
                try:
                    if obj.status.s == 'Перенесена':
                        message = '{}, здравствуйте!\nВаш отчет не был сформирован так как ваш доход не достиг порог оплаты ($100). Оплата перенесена.\nПериод: {},{}\nСумма: {} \nСтатус: Перенесена'.format(fio, month, str(year), str(summa))
                        doc_send = False
                except:
                    fewrfwef = 0
                token = TOKEN

                bot = telegram.Bot(token=token)
                i_accounts = InlineKeyboardButton(text='Отчеты', callback_data='accounts')
                i_setting = InlineKeyboardButton(text='Настройки', callback_data='setting')
                mrk = InlineKeyboardMarkup([[i_accounts, i_setting]])
                bot.sendMessage(chat_id=user_id, text=message)
                if doc_send:
                    bot.sendDocument(chat_id=user_id, document=obj.document)
                bot.sendMessage(chat_id=user_id, text='Вы можете просмотреть свои отчеты нажав на кнопку Отчеты или поменять настройки вашего аккаунта, нажав на кнопку Настройки', reply_markup=mrk)
            except:
                dwsdws = 0
            #_end_bot


        stories.objects.create(obj_id=obj.pk, admin=self.request.user, obj='Отчет', text="Создан")

        
    
        try:
            status = obj.status.s
        except:
            status = ''
        pre = acc_pre_value.objects.create(pk=obj.pk, pseudonym = obj.pseudonym, month = obj.month, year = obj.year, status = status, day_payment = obj.day_payment, document = obj.document, summa = obj.summa, card_number = obj.card_number)

        return context


class AccUpdateView(LoginRequiredMixin, DetailView):
    model = Account
    
    def get_context_data(self, *args, **kwargs):
        context=super().get_context_data(*args, **kwargs)
        obj = context['object']
        
        year = obj.year
        summa = obj.summa
        mo = obj.month.id
        months = ['Январь', 'Февраль', 'Март', 'Апрель', 'Май', 'Июнь', 'Июль', 'Август', 'Сентябрь', 'Октябрь', 'Ноябрь', 'Декабрь']
        month = months[mo-1]
        #__bot
        if obj.notification == 'Вкл':
            try:
                prof = Profile.objects.get(pseudonym=obj.pseudonym)
                fio = prof.name
                pre = acc_pre_value.objects.filter(pk=obj.pk)[0]
                if obj.status.s == 'Оплачен' and pre.status != 'Оплачен':
                    message = 'Здравствуйте,{}\n  Мы оплатили вашего дохода за {} {}.'.format(fio, month, str(year))
                    doc_send = True

                if obj.status.s == 'Перенесена' and pre.status != 'Перенесена':
                    message = '{}, здравствуйте!\nВаш отчет не был сформирован так как ваш доход не достиг порог оплаты ($100). Оплата перенесена.\nПериод: {},{}\nСумма: {} \nСтатус: Перенесена'.format(fio, month, str(year), str(summa))
                    doc_send = False
                l = prof.login
                user_id = subscribersbot.objects.get(login = l).user_id


                token = TOKEN
                bot = telegram.Bot(token=token)
                i_accounts = InlineKeyboardButton(text='Отчеты', callback_data='accounts')
                i_setting = InlineKeyboardButton(text='Настройки', callback_data='setting')
                mrk = InlineKeyboardMarkup([[i_accounts, i_setting]])
                bot.sendMessage(chat_id=user_id, text=message)
                if doc_send:
                    bot.sendDocument(chat_id=user_id, document=obj.document)
                bot.sendMessage(chat_id=user_id, text='Вы можете просмотреть свои отчеты нажав на кнопку Отчеты или поменять настройки вашего аккаунта, нажав на кнопку Настройки', reply_markup=mrk)
            except:
                dwsdws = 0
            #_end_bot

        
        now = Account.objects.filter(pk=obj.pk).values()[0]
        c = 0
        try:
            pre = acc_pre_value.objects.filter(pk=obj.pk).values()[0]
            text = ''
            for i in pre:

                if i == 'id':
                    feef = 0
                elif i == 'published':
                    wefw = 0
                elif i == 'comment':
                    fwf = 0
                elif i == 'notification':
                    dwdwe = 0
                elif i == 'status':
                    try:
                        if pre[i] != obj.status.s:
                            text+=i + ':\n' + pre[i] + ' >> ' + obj.status.s + ';    '
                    except:
                        if pre[i] != '':
                            text+=i + ':\n' + pre[i] + ' >> ' + "-------" + ';     '
                
                elif i == 'month':
                    
                    months = ['Январь', 'Февраль', 'Март', 'Апрель', 'Май', 'Июнь', 'Июль', 'Август', 'Сентябрь', 'Октябрь', 'Ноябрь', 'Декабрь']
                    month = months[now['month_id']-1]
                    if pre[i] != month:         
                        text+=i + ':\n' + pre[i] + ' >> ' + month + ';    '
                elif pre[i] != now[i]:
                    
                    text+=i + ':\n' + str(pre[i]) + ' >> ' + str(now[i]) + ';    '
            if text:
                stories.objects.create(obj_id=obj.pk, admin=self.request.user, obj='Отчет', text=text)
                
                
            
        except:
            dedwef = 0

            
        # save currently fields as old , that next time will be pre
        try:
            
            pre = acc_pre_value.objects.get(pk=obj.pk)
            if pre.pseudonym != obj.pseudonym:            
                pre.pseudonym = obj.pseudonym

            if str(pre.month) != str(obj.month):
                pre.month = str(obj.month)

            if pre.year != obj.year:    
                pre.year = (obj.year)

            try:
                if pre.status != obj.status.s:
                    pre.status = obj.status.s

            except:
                pre.status = ''

            if pre.day_payment != obj.day_payment:
                pre.day_payment = obj.day_payment


            if pre.document != obj.document:    
                pre.document = obj.document

        
            if pre.summa != obj.summa:
                pre.summa = obj.summa

            
            if pre.card_number != obj.card_number:
                pre.card_number = obj.card_number


            pre.save()
            
        except:
            try:
                status = obj.status.s
            except:
                status = ''
            pre = acc_pre_value.objects.create(pk=obj.pk, pseudonym = obj.pseudonym, month = obj.month, year = obj.year, status = status, day_payment = obj.day_payment, document = obj.document, summa = obj.summa, card_number = obj.card_number)

        return context



class ProfCreateView(LoginRequiredMixin, CreateView, PermissionRequiredMixin):
    template_name = 'bot/createprof.html'
    form_class = ProfForm
    success_url = '/new/{id}'
    permanent = True
    permission_required = 'hello.add_profile'
    def get_context_data(self, *args, **kwargs):
        con = super().get_context_data(*args, **kwargs)
        con['profiles'] = Profile.objects.all()
        return con

class ProfDeleteView(LoginRequiredMixin, DeleteView, PermissionRequiredMixin):
    model = Profile
    permission_required = 'hello.delete_profile'
    def get_success_url(self, *args, **kwargs):
        con = super().get_context_data(*args, **kwargs)
        obj = con['object']
        ps = obj.pseudonym
        login = obj.login


        url = '/afterdeleting/{}/{}/'.format(ps, login)
        return url
@login_required
def afterdeleting(request, ps, login):
    for i in Account.objects.filter(pseudonym=ps):
        i.delete()
    try:
        s = subscribersbot.objects.get(login=login)
        user_id = s.user_id
        s.delete()
        ch = changing.objects.get(user_id=user_id)
        ch.delete()
        t = typing.objects.get(user_id=user_id)
        t.delete()
        requests.get('https://api.telegram.org/bot{}/sendMessage?chat_id={}&text=ваш профиль удален\nнажмите /start, чтобы войти снова'.format(TOKEN, user_id))
    except:
        fwrfwe = 0
    return redirect(Profiles)


class AccEditView(LoginRequiredMixin, UpdateView, PermissionRequiredMixin):
    model = Account
    form_class = AccForm
    success_url = '/accupdate/{id}'
    permission_required = 'hello.change_account'
    def get_context_data(self, *args, **kwargs):
        con = super().get_context_data(*args, **kwargs)
        con['profiles'] = Profile.objects.all().order_by('pseudonym')
        months = ['Январь', 'Февраль', 'Март', 'Апрель', 'Май', 'Июнь', 'Июль', 'Август', 'Сентябрь', 'Октябрь', 'Ноябрь', 'Декабрь']
        con['months'] = months
        obj = con['object']
        con['stories'] = stories.objects.filter(obj_id=obj.pk, obj='Отчет').order_by('-published')
        card_list = []
        for i in card.objects.filter(pseudonym=obj.pseudonym):
            card_list.append(i.card_number)
        
        con['card'] = card_list
        return con
@login_required
def Profiles(request):
    p = os.listdir(os.path.join(BASE_DIR, 'files/contract/main'))
    if len(p) == 2:
        search_dir = os.path.join(BASE_DIR, 'files/contract/main')
        os.chdir(search_dir)
        files = filter(os.path.isfile, os.listdir(search_dir))
        files = [os.path.join(search_dir, f) for f in files] # add path to each file
        files.sort(key=lambda x: os.path.getmtime(x))
        os.remove(files[0])


    profiles = Profile.objects.all().order_by('prefix')
    bot_users = []
    for i in subscribersbot.objects.all():
        bot_users.append(i.login)
    context = {"profiles": profiles, 'ps': 'Все', 'allprofiles': profiles, 'path': 'Все', 'by': 'pseudonym', 'bot_users': bot_users}
    return render(request, 'bot/profiles.html', context)


@login_required
def SortProfiles(request, ps, by):
    allprofiles = Profile.objects.all()
    if ps == 'Все':
        profiles = Profile.objects.all().order_by('pseudonym').order_by(by)
    else:
        profiles = Profile.objects.filter(pseudonym=ps)
    bot_users = []
    for i in subscribersbot.objects.all():
        bot_users.append(i.login)
    context = {"profiles": profiles, 'ps': ps, 'allprofiles': allprofiles, 'path': ps, 'by': by, 'bot_users': bot_users}
    return render(request, 'bot/profiles.html', context)    

#______bot:
def bot(request):
    
    
    
    
    
    return redirect(open_site)
@login_required
def open_site(request):
    if request.method == 'POST':
        f = SecurityForm(request.POST)
        
        if f.is_valid():
            
            try:
                parol = security.objects.get(pk=1).parol
                if f.cleaned_data['parol'] == parol:
                    return redirect('folder')
                else:
                    return render(request, 'bot/security.html', {'form': f, 'message': 'неверно, пожалуйста, введите пароль еще раз'})
            except:
                
                security.objects.create(parol = f.cleaned_data['parol'])
                return redirect('folder')
        else:
            return render(request, 'bot/security.html', {'form': f, 'message':'Введите пароль'})
    else:
        f = SecurityForm()
        return render(request, 'bot/security.html', {'form': f, 'message':'Введите пароль'})



class SecEditView(UpdateView, LoginRequiredMixin):
    model = security
    form_class = SecurityForm
    success_url = '/'


@login_required
def addadmin(request):
    if request.method == 'POST':
        f = Addadmin(request.POST)
        if f.is_valid():
            obj = User.objects.create_user(username=f.cleaned_data['username'], password=f.cleaned_data['password'], email=f.cleaned_data['email'])
            title = f.cleaned_data['title']
            if title == 'Бухгалтерия':
                exp = ['Can delete account', 'Can delete profile']
            elif title == 'Менеджер':
                exp = ['Can delete account', 'Can delete profile', 'Can add account', 'Can change account']
            elif title == 'Аналитик':
                exp = ['Can delete account', 'Can delete profile', 'Can add account', 'Can change account', 'Can change profile', 'Can add profile']
            permissions = Permission.objects.all()
            for p in permissions:
                if not p.name in exp:
                    obj.user_permissions.add(p)
            obj.save()
            return redirect('editpassword')
        else:
            context = {'form': f}
            return render(request, 'bot/addadmin.html', context)
    else:
        f = Addadmin
        context = {'form': f}
        return render(request, 'bot/addadmin.html', context)



@login_required
def editadmin(request, pk):
    if request.method == 'POST':
        f = Addadmin(request.POST)
        if f.is_valid():

            username=f.cleaned_data['username']
            password=f.cleaned_data['password']
            email=f.cleaned_data['email']
            title = f.cleaned_data['title']
            if password != '':
                obj = User.objects.get(pk=pk)
                obj.delete()
                obj = User.objects.create_user(username=f.cleaned_data['username'], password=f.cleaned_data['password'], email=f.cleaned_data['email'])
            else:
                obj = User.objects.get(pk=pk)
                obj.username=f.cleaned_data['username']
                obj.email=f.cleaned_data['email']
            if title == 'Бухгалтерия':
                exp = ['Can delete account', 'Can delete profile']
            elif title == 'Менеджер':
                exp = ['Can delete account', 'Can delete profile', 'Can add account', 'Can change account']
            elif title == 'Аналитик':
                exp = ['Can delete account', 'Can delete profile', 'Can add account', 'Can change account', 'Can change profile', 'Can add profile']
            permissions = Permission.objects.all()
            for p in permissions:
                try:
                    obj.user_permissions.remove(p)
                except:
                    do = 0
            for p in permissions:
                if not p.name in exp:
                    obj.user_permissions.add(p)

            obj.save()
            return redirect('editpassword')
        else:
            obj = User.objects.get(pk=pk)
            username=obj.username
            password=obj.password
            email=obj.email
            if obj.user_permissions.filter(name='Can add account'):
                title = 'Бухгалтерия'
            
            elif obj.user_permissions.filter(name='Can add profile'):
                title = 'Менеджер'
            else:
                title = 'Аналитик'
    
            context = {'form': f, 'username': username, 'password': password, 'email': email, 'title': title}
            return render(request, 'bot/editadmin.html', context)
    else:
        f = Addadmin
    
        obj = User.objects.get(pk=pk)
        username=obj.username
        password=obj.password
        email=obj.email
        if obj.user_permissions.filter(name='Can add account'):
            title = 'Бухгалтерия'
        
        elif obj.user_permissions.filter(name='Can add profile'):
            title = 'Менеджер'
        else:
            title = 'Аналитик'
        context = {'form': f, 'username': username, 'password': password, 'email': email, 'title': title}
        
        return render(request, 'bot/editadmin.html', context)

            
@login_required
def sendfile(request, y, m, d, f):

    a = Account.objects.filter(document='{}/{}/{}/{}'.format(y, m, d, f))
    return FileResponse(a[0].document)

def send_photo_ava(request, f):
    obj = Profile.objects.filter(photo='profile_photos/{}'.format(f))
    return FileResponse(obj[0].photo)


@login_required
def senddocument(request, f):
    a = Profile.objects.filter(passport_main_st='{}'.format(f))
    if not a:
        a = Profile.objects.filter(passport_st_registration='{}'.format(f))
        
        if not a:
            a = Profile.objects.filter(signed_contract='{}'.format(f))
            return FileResponse(a[0].signed_contract)
        else:
            return FileResponse(a[0].passport_st_registration)
    else:
        return FileResponse(a[0].passport_main_st)


@login_required
def admins(request):
    users = User.objects.all()

    title = []
    for i in users:
        if i.user_permissions.filter(name='Can add account'):
            title.append('Бухгалтерия')
        elif i.user_permissions.filter(name='Can add profile'):
            title.append('Менеджер')
        else:
            title.append('Аналитик')
    context = {'users': users, 'title': title}
    return render(request, 'bot/admins.html', context)

@login_required
def deladmin(request, pk):
    try:
        deluser = User.objects.get(pk=pk)
        deluser.delete()
    except:
        refer=0
    users = User.objects.all()
    title = []        
    for i in users:
        if i.user_permissions.filter(name='Can add account'):
            title.append('Бухгалтерия')
        elif i.user_permissions.filter(name='Can add profile'):
            title.append('Менеджер')
        else:
            title.append('Аналитик')
    
    context = {'users': users, 'title': title}
    return render(request, 'bot/admins.html', context)

class HappybirthdayEditView(UpdateView, LoginRequiredMixin):  #
    model = happybirthday
    form_class = HappybirthdayForm
    success_url = '/calendar'


@login_required
def stories_admin(request, username):
    story = stories.objects.filter(admin=username)
    users = User.objects.all()
    context = {'story': story, 'users': users, 'username': username}
    return render(request, 'bot/stories_admin.html', context)


@login_required
def calendar(request):
    all_list = {'Январь': [], 'Февраль': [], 'Март': [], 'Апрель': [], 'Май': [], 'Июнь': [], 'Июль': [], 'Август': [], 'Сентябрь': [], 'Октябрь': [], 'Ноябрь': [], 'Декабрь': []}
    n_month = 1
    days_of_the_months = {'Январь': 31, 'Февраль': 28, 'Март': 31, 'Апрель': 30, 'Май': 31, 'Июнь': 30, 'Июль': 31, 'Август': 31, 'Сентябрь':30, 'Октябрь': 31, 'Ноябрь': 30, 'Декабрь': 31}
    if int(datetime.now().year) % 4 == 0:
        days_of_the_months['Февраль'] = 29
    for i in days_of_the_months:
        for d in range(1, days_of_the_months[i]+1):
            obj = Profile.objects.filter(date_birthday__icontains='-{}-'.format(str(n_month))).filter(date_birthday__endswith='-'+str(d))
            if obj:
                ps = ''
                for user in obj:
                    ps += user.pseudonym + ',\n'
                all_list[i].append(ps)
            else:
                all_list[i].append('None')


        
        n_month+=1
    for i in all_list:
        for p in range(5):
            all_list[i].append('Null')
    context = {'all_list': all_list}
    return render(request, 'bot/calendar.html', context)



@login_required
def Contract(request, ps):
    obj = Profile.objects.get(pseudonym=ps)
    p = os.listdir(os.path.join(BASE_DIR, 'files/contract/main'))
    document = Document(os.path.join(BASE_DIR, 'files/contract/main/'+p[0]))
    for table in document.tables:
        for r in table.rows:
            for c in r.cells:
                for p in c.paragraphs:

                    try:
                        if obj.name != None:
                            p.text = p.text.replace('{Ф.И.О. партнера}', obj.name)

                        else:
                            p.text = p.text.replace('{Ф.И.О. партнера}', '_______')
                        if obj.prefix != None:
                            p.text = p.text.replace('{номер договора}', obj.prefix)
                        else:
                            p.text = p.text.replace('{номер договора}', '__________')
                        y, m, d = obj.published.split('-')
                        p.text = p.text.replace('{день}', d)
                        p.text = p.text.replace('{месяц}', m)
                        p.text = p.text.replace('{год}', y)
                        if obj.reward != None:
                            p.text = p.text.replace('{вознаграждение%}', obj.reward)
                        else:
                            p.text = p.text.replace('{вознаграждение%}', '__________')
                        p.text = p.text.replace('{(вознаграждение словами)}', 'вознаграждение')
                        if obj.email != None:
                            p.text = p.text.replace('{электронная почта партнера}', obj.email)
                        else:
                            p.text = p.text.replace('{электронная почта партнера}', '_________')
                        if obj.number_passport != None:

                            p.text = p.text.replace('{Серия и номер}', obj.number_passport)
                        else:
                            p.text = p.text.replace('{Серия и номер}', '________')
                        y, m, d = obj.date_passport.split('-')
                        if obj.date_passport != None:

                            p.text = p.text.replace('{дата выдачи в формате дд.мм.гггг}', '{}.{}.{}'.format(d, m, y))
                        else:
                            
                            p.text = p.text.replace('{дата выдачи в формате дд.мм.гггг}', '________')
                        if obj.who_gave != None:
                            
                            p.text = p.text.replace('{кем выдан}', obj.who_gave)
                        else:
                            
                            p.text = p.text.replace('{кем выдан}', '________')
                        if obj.inn != None:
                           
                            p.text = p.text.replace('{инн}', obj.inn)
                        else:
                            
                            p.text = p.text.replace('{инн}', '______')
                        if obj.card_number != None:
                            p.text = p.text.replace('{номер карты}', obj.card_number)
                        else:
                            
                            p.text = p.text.replace('{номер карты}', '_______')
                        if obj.bank != None:

                            p.text = p.text.replace('{банк}', obj.bank)
                        else:
                            
                            p.text = p.text.replace('{банк}', '_________')
                        if obj.name != None:

                            p.text = p.text.replace('{фио}', obj.name)
                        else:
                            p.text = p.text.replace('{фио}', '______')
                    except:
                        dedede = 0
    pse = obj.pseudonym
    pse = pse.replace("'", "")
    pse = pse.replace('&', '')
    pse = pse.replace('(', '')
    pse = pse.replace(')', '')
    document.save('files/contract/{}.docx'.format(pse))

    p = os.path.abspath('files/contract/{}.docx'.format(pse))
    p=p.replace(' ', '\ ')
    os.system('unoconv -f pdf {}'.format(p))  # sudo apt install unoconv python3-unoconv;    copy uno.py and unohelper.py to python3.8 or more from python3 
    f = open('files/contract/{}.pdf'.format(pse), 'rb')
    return FileResponse(f)

@login_required
def get_contract_main(request, file):
    f = open('files/contract/main/{}'.format(file))
    return FileResponse(f)

class ContractEditView(UpdateView):
    model = contract
    form_class = ContractForm
    success_url = '/profiles'






















def audio_create(request, artist, extra):
    FormSet = modelformset_factory(Audio, extra=extra, fields=('composition', 'artist', 'autor_music', 'autor_text', 'album', 'isrc', 'genre', 'copyright', 'related_rights', 'upc', 'release_date', 'territory', 'link'),
        labels = {
            'composition': '', 
            'artist': '', 
            'autor_music': '', 
            'autor_text': '', 
            'album': '', 
            'isrc': '', 
            'genre': '', 
            'copyright': '', 
            'related_rights': '', 
            'upc': '', 
            'release_date': '', 
            'territory': '', 
            'link': '',
        },
        widgets = {
            'composition': forms.TextInput(attrs={'class': 'composition'}),
            'artist': forms.TextInput(attrs={'class': 'artist'}),
            'autor_music': forms.TextInput(attrs={'class': 'autor_music'}),
            'autor_text': forms.TextInput(attrs={'class': 'autor_text'}),
            'album': forms.TextInput(attrs={'class': 'album'}),
            'isrc': forms.TextInput(attrs={'class': 'isrc'}),
            'genre': forms.TextInput(attrs={'class': 'genre'}),
            'copyright': forms.TextInput(attrs={'class': 'copyright'}),
            'related_rights': forms.TextInput(attrs={'class': 'related_rights'}),
            'upc': forms.TextInput(attrs={'class': 'upc'}),
            'release_date': forms.TextInput(attrs={'class': 'release_date'}),
            'territory': forms.TextInput(attrs={'class': 'territory'}),
            'link': forms.TextInput(attrs={'class': 'link'}),
        }

    )
    if request.method == 'POST':
        formset = FormSet(request.POST, queryset=Audio.objects.filter(pseudonym=artist))
        if formset.is_valid():
            formset.save()
            for form in formset:
                if form.instance.pseudonym == None and form.instance.artist != None:
                    form.instance.pseudonym = artist
                    form.instance.status = '+'
                    current_pk = form.instance.pk
                    try:
                        video_pk = max(list(Video.objects.all().values_list('pk')))[0]
                        if current_pk > video_pk:
                            do_nothing=True
                        else:
                            form.instance.pk = video_pk + 1
                    except:
                        dwerdfwerf=21
                    form.save()

        return redirect(content, ps=artist)

    else:
        formset = FormSet(queryset=Audio.objects.filter(pseudonym=None))
        
        context = {'form': formset, 'add': extra+1, 'artist': artist}
        return render(request, 'bot/createaudio.html', context)


class AudioCreateView(LoginRequiredMixin, CreateView):
    template_name = 'bot/createaudio.html'
    form_class = AudioForm
    success_url = '/audio_detail/{id}'
    def get_context_data(self, *args, **kwargs):
        context = super().get_context_data(*args, **kwargs)
        ps = []
        for i in Profile.objects.all().order_by('pseudonym'):
            ps.append(i.pseudonym)
        context['profiles'] = ps
        context['allprofiles'] = Profile.objects.all().order_by('pseudonym')
        return context

class AuidoEditView(LoginRequiredMixin, UpdateView):
    model = Audio
    form_class = AudioForm
    def get_success_url(self, *args, **kwargs):
        context = super().get_context_data(*args, **kwargs)
        ps = context['object'].pseudonym
        return '/content/{}'.format(ps)

class AudioDetailView(LoginRequiredMixin, DetailView):
    model = Audio
@login_required
def delete_audio(request, pk):
    audio = Audio.objects.get(pk=pk)
    artist = audio.pseudonym
    if audio.status == '+':
        audio.status = '-'
    else:
        audio.status = '+'
    audio.save()
    return redirect(content, ps=artist)




def video_create(request, artist, extra):
    FormSet = modelformset_factory(Video, extra=extra, fields=('composition', 'type', 'artist', 'isrc', 'producer', 'operator', 'autor_script', 'painter', 'copyright', 'release_date', 'territory', 'link'),
        labels = {
            'composition': '', 
            'type': '', 
            'artist': '', 
            'isrc': '', 
            'producer': '', 
            'operator': '', 
            'autor_script': '', 
            'painter': '', 
            'copyright': '', 
            'release_date': '', 
            'territory': '', 
            'link': '',
        },
        widgets = {
            'composition': forms.TextInput(attrs={'class': 'composition'}), 
            'type': forms.TextInput(attrs={'class': 'type'}), 
            'artist': forms.TextInput(attrs={'class': 'artist'}), 
            'isrc': forms.TextInput(attrs={'class': 'isrc'}), 
            'producer': forms.TextInput(attrs={'class': 'producer'}), 
            'operator': forms.TextInput(attrs={'class': 'operator'}), 
            'autor_script': forms.TextInput(attrs={'class': 'autor_script'}), 
            'painter': forms.TextInput(attrs={'class': 'painter'}), 
            'copyright': forms.TextInput(attrs={'class': 'copyright'}), 
            'release_date': forms.TextInput(attrs={'class': 'release_date'}),
            'territory': forms.TextInput(attrs={'class': 'territory'}), 
            'link': forms.TextInput(attrs={'class': 'link'}),
        }

    )
    if request.method == 'POST':
        formset = FormSet(request.POST, queryset=Video.objects.filter(pseudonym=artist))
        if formset.is_valid():
            formset.save()
            for form in formset:
                if form.instance.pseudonym == None and form.instance.artist != None:
                    form.instance.pseudonym = artist
                    form.instance.status = '+'
                    current_pk = form.instance.pk
                    try:
                        audio_pk = max(list(Audio.objects.all().values_list('pk')))[0]

                        if current_pk > audio_pk:
                            do_nothing=True
                        else:
                            form.instance.pk = audio_pk + 1
                    except:
                        dedwe=91
                    form.save()
        return redirect(content, ps=artist)

    else:
        formset = FormSet(queryset=Video.objects.filter(pseudonym=None))
        
        context = {'form': formset, 'add': extra+1, 'artist': artist}
        return render(request, 'bot/createvideo.html', context)



class VideoCreateView(LoginRequiredMixin, CreateView):
    template_name = 'bot/createvideo.html'
    form_class = VideoForm
    success_url = '/video_detail/{id}'
    def get_context_data(self, *args, **kwargs):
        context = super().get_context_data(*args, **kwargs)
        ps = []
        for i in Profile.objects.all().order_by('pseudonym'):
            ps.append(i.pseudonym)
        context['profiles'] = ps
        context['allprofiles'] = Profile.objects.all().order_by('pseudonym')
        return context


class VideoEditView(LoginRequiredMixin, UpdateView):
    model = Video
    form_class = VideoForm
    def get_success_url(self, *args, **kwargs):
        context = super().get_context_data(*args, **kwargs)
        ps = context['object'].pseudonym
        return '/content/{}'.format(ps)


class VideoDetailView(LoginRequiredMixin, DetailView):
    model = Video
@login_required
def delete_video(request, pk):
    video = Video.objects.get(pk=pk)
    artist = video.pseudonym
    if video.status == '+':
        video.status = '-'
    else:
        video.status = '+'
    video.save()
    return redirect(content, ps=artist)
@login_required
def content(request, ps):
    audios = Audio.objects.filter(pseudonym=ps)
    indexes = {}

    for a in audios:
        obj = Audio.objects.filter(pseudonym=a.pseudonym, composition=a.composition, artist=a.artist, autor_music=a.autor_music, autor_text=a.autor_text, album=a.album, isrc=a.isrc, genre=a.genre, copyright=a.copyright, related_rights=a.related_rights, upc=a.upc, release_date=a.release_date, territory=a.territory)
        obj_pks = [i.pk for i in obj]
        index = obj_pks.index(a.pk) + 1
        indexes[a.pk] = index

    videos = Video.objects.filter(pseudonym=ps)
    indexes_video = {}

    for v in videos:
        obj = Video.objects.filter(pseudonym=v.pseudonym, composition=v.composition, artist=v.artist, isrc=v.isrc, copyright=v.copyright, release_date=v.release_date, territory=v.territory, type=v.type, producer=v.producer, operator=v.operator, autor_script=v.autor_script, painter=v.painter)
        obj_pks = [i.pk for i in obj]
        index = obj_pks.index(v.pk) + 1
        indexes_video[v.pk] = index

    context = {'audios': audios, 'videos': videos, 'ps': ps, 'indexes': indexes, 'indexes_video': indexes_video}
    return render(request, 'bot/content.html', context)


@login_required
def generation_file(request, artist):
    file_path = os.path.join(BASE_DIR, 'files/app/main shablon.docx')
    folder_path = os.path.join(BASE_DIR, 'files/app/')
    
    list_audio = list(Audio.objects.filter(pseudonym=artist).values_list())
    list_video = list(Video.objects.filter(pseudonym=artist).values_list())
    
    doc = Document(file_path)
    
    n = 1

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    def modifyBorder(table):  # set border size
        tbl = table._tbl # get xml element in table
        for cell in tbl.iter_tcs():
            tcPr = cell.tcPr # get tcPr element, in which we can define style of borders
            tcBorders = OxmlElement('w:tcBorders')
            top = OxmlElement('w:top')
            top.set(qn('w:val'), 'nil')

            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'nil')

            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'nil')
            bottom.set(qn('w:sz'), '30')
            bottom.set(qn('w:space'), '0')
            bottom.set(qn('w:color'), 'black')

            right = OxmlElement('w:right')
            right.set(qn('w:val'), 'nil')

            tcBorders.append(top)
            tcBorders.append(left)
            tcBorders.append(bottom)
            tcBorders.append(right)
            tcPr.append(tcBorders)
    for table in doc.tables:
        if n == 2:
            modifyBorder(table)
            index = 1
            for l in list_audio:
                if l[-3] == '-':
                    continue
                if l[-2] == False:
                    continue
                rows = table.add_row()
                cn = 0
                for c in rows.cells:
                    if cn == 0:
                        c.text = str(index)
                    else:
                        c.text = l[cn]
                    cn += 1
                    p = c.paragraphs[0]    
                    for run in p.runs:
                        run.font.size = Pt(10)
                number = str(len(os.listdir(os.path.join(BASE_DIR, 'files/app/{}'.format(artist)))) + 1)
                obj = Audio.objects.get(pk=int(l[0]))
                obj.number_of_file = number
                obj.save()
                index += 1


        elif n == 3:
            modifyBorder(table)
            index = 1
            for l in list_video:
                if l[-3] == '-':
                    continue
                if l[-2] == False:
                    continue
                rows = table.add_row()
                cn = 0
                for c in rows.cells:
                    if cn == 0:
                        c.text = str(index)
                    else:
                        c.text = l[cn]
                    cn += 1
                    p = c.paragraphs[0]    
                    for run in p.runs:
                        run.font.size = Pt(10)
                number = str(len(os.listdir(os.path.join(BASE_DIR, 'files/app/{}'.format(artist)))) + 1)
                obj = Video.objects.get(pk=int(l[0]))
                obj.number_of_file = number
                obj.save()
                index += 1
       
        else:
            artist = artist.replace('_', ' ')
            obj = Profile.objects.get(pseudonym=artist)
            for r in table.rows:
                for c in r.cells:
                    for p in c.paragraphs:
                        if obj.name != None:
                            p.text = p.text.replace('{name}', obj.name)
                        else:
                            p.text = p.text.replace('{name}', '________')
                        if obj.prefix != None:
                            p.text = p.text.replace('{prefix}', obj.prefix)
                        else:
                            p.text = p.text.replace('{prefix}', '_____')
                        artist = artist.replace(' ', '_')
                        try:
                            index = len(os.listdir(folder_path+artist))+1
                        except:
                            index = 1
                        p.text = p.text.replace('{index}', str(index))
        
                        y, m, d = str(obj.published).split('-')
                        p.text = p.text.replace('{date}', '{}.{}.{}'.format(d, m, y))
                        months = ['Январь', 'Февраль', 'Март', 'Апрель', 'Май', 'Июнь', 'Июль', 'Август', 'Сентябрь', 'Октябрь', 'Ноябрь', 'Декабрь']
                        p.text = p.text.replace('{year}', y)
                        p.text = p.text.replace('{month}', months[int(m)-1])
                        p.text = p.text.replace('{day}', d)
                        

        n += 1
    index = str(len(os.listdir(os.path.join(BASE_DIR, 'files/app/{}'.format(artist)))) + 1)
    try:
        d = datetime.now()
        date_time = '__{}-{}-{}-{}-{}-{}'.format(str(d.year), str(d.month), str(d.day), str(d.hour), str(d.minute), str(d.second))
        artist = artist.replace(' ', '_')
        doc.save(folder_path+artist+'/'+index+'.'+artist+date_time+'.docx')
    except:
        artist = artist.replace(' ', '_')
        os.system('mkdir {}'.format(folder_path+artist))
        doc.save(folder_path+artist+'/'+index+'.'+artist+date_time+'.docx')
    f = open(folder_path+artist+'/'+index+'.'+artist+date_time+'.docx', 'rb')
    for i in Audio.objects.filter(pseudonym=artist):
        i.is_generate=False
        i.save()
    for i in Video.objects.filter(pseudonym=artist):
        i.is_generate=False
        i.save()
    return FileResponse(f)

@login_required
def app_list(request, artist):
    try:
        artist = artist.replace(' ', '_')
        search_dir = os.path.join(BASE_DIR, 'files/app/{}'.format(artist))
        all_files = os.listdir(os.path.join(BASE_DIR, 'files/app/{}'.format(artist)))
        all_files = [os.path.join(search_dir, f) for f in all_files]
        all_files.sort(key=lambda x: os.path.getmtime(x))  #sort byd date
        all_apps = []
        # removing path name, rest only file name
        for i in all_files:
            *path, file = i.split('/')
            all_apps.append(file)

        # all_apps.sort(key=os.path.getctime)
        date_time = []
        for i in all_apps:
            folder, d_t_docx = i.split('__')
            d_t, docx = d_t_docx.split('.')
            year, month, day, hour, minute, second = d_t.split('-')
            date_time.append('{}.{}.{} {}:{}'.format(day, month, year, hour, minute))
        context = {'all_apps': all_apps, 'date_time': date_time}
    except:
        context = {'all_apps': [], 'date_time': []}
    return render(request, 'bot/app_list.html', context)

@login_required
def open_app(request, app):

    number_folder, date = app.split('__')
    n, folder = number_folder.split('.')
    file_path = os.path.join(BASE_DIR, 'files/app/')
    f = open(file_path+folder+'/'+app, 'rb')
    return FileResponse(f)


def change_generate_audio(request, pk):
    obj = Audio.objects.get(pk=pk)
    if obj.is_generate:
        obj.is_generate = False
    else:
        obj.is_generate = True
    obj.save()
    return redirect(content, ps=obj.pseudonym)

def change_generate_video(request, pk):
    obj = Video.objects.get(pk=pk)
    if obj.is_generate:
        obj.is_generate = False
    else:
        obj.is_generate = True
    obj.save()
    return redirect(content, ps=obj.pseudonym)
    